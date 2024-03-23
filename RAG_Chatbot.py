import streamlit as st
from streamlit_chat import message
from langchain.chains import ConversationalRetrievalChain
import requests
from bs4 import BeautifulSoup
import base64
import docx
import os
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Pinecone
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate, ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from dotenv import load_dotenv
from pinecone import Pinecone as pc
from langchain.memory import ConversationBufferMemory
from langchain_core.runnables import RunnableParallel, RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser

load_dotenv()
os.getenv("OPENAI_API_KEY")
index_name=os.getenv("PINECONE_INDEX_NAME")

pc(
    api_key=os.getenv("PINECONE_API_KEY"),
    environment=os.getenv("PINECONE_ENV_NAME")
    )

class Document:
    def __init__(self, content, metadata=None):
        self.page_content = content
        self.metadata = metadata if metadata is not None else {}

def initialize_session_state():
    if 'history' not in st.session_state:
        st.session_state['history'] = []

    if 'generated' not in st.session_state:
        st.session_state['generated'] = ["Hello! Ask me anything about 🤗"]

    if 'past' not in st.session_state:
        st.session_state['past'] = ["Hey! 👋"]

def conversation_chat(query, chain, history):
    result = chain({"question": query, "chat_history": history})
    history.append((query, result["answer"]))
    return result["answer"]

def display_chat_history(chain):
    reply_container = st.container()
    container = st.container()

    with container:
        with st.form(key='my_form', clear_on_submit=True):
            user_input = st.text_input("Question:", placeholder="Ask about your Documents", key='input')
            submit_button = st.form_submit_button(label='Send')

        if submit_button and user_input:
            with st.spinner('Generating response...'):
                output = conversation_chat(user_input, chain, st.session_state['history'])

            st.session_state['past'].append(user_input)
            st.session_state['generated'].append(output)

    if st.session_state['generated']:
        with reply_container:
            for i in range(len(st.session_state['generated'])):
                message(st.session_state["past"][i], is_user=True, key=str(i) + '_user', avatar_style="thumbs")
                message(st.session_state["generated"][i], key=str(i), avatar_style="fun-emoji")


def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def read_md_or_txt(file):
    return file.getvalue().decode()

def scrape_webpage(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # This will raise an HTTPError if the HTTP request returned an unsuccessful status code.

        soup = BeautifulSoup(response.text, 'html.parser')
        return soup.get_text(separator='\n')  # Using separator for better text formatting
    except requests.RequestException as e:
        return f"Error fetching URL: {e}"
    except Exception as e:
        return f"An error occurred: {e}"


def get_pdf_text(pdf_docs):
    text=""
    for pdf in pdf_docs:
        pdf_reader= PdfReader(pdf)
        for page in pdf_reader.pages:
            text+= page.extract_text()
    return  text

def process_uploaded_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        return get_pdf_text([uploaded_file])
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx(uploaded_file)
    else:  # Assuming text or markdown
        return read_md_or_txt(uploaded_file)
    return text

def get_text_chunks(text, metadata):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    chunks = text_splitter.split_text(text)
    return [Document(chunk, metadata) for chunk in chunks]

def init_embeddings():
    embeddings = OpenAIEmbeddings()
    return embeddings

def get_vector_store(text_chunks):
    
    vector_store = Pinecone.from_existing_index(index_name, init_embeddings())
    vector_store.add_documents(text_chunks)


def get_conversational_chain(vector_store,user_question,chat_history):
    # Construct the conversational retrieval chain
    model = ChatOpenAI(model_name='gpt-3.5-turbo', temperature=0, openai_api_key=os.environ["OPENAI_API_KEY"])
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    
    #Define Prompts
    system_template = r'''
    Use the following pieces of context to answer the user's question.
    If you find the answer in the provided context, start your response with "Thank you for your question. 
    According to my knowledge,.." and then give a brief summary of the response using short sentences.
    If you don't find the answer in the provided context, just respond "I dont have the response in this context. Can you please try asking again?"
    If the question is written in gibberish words and cannot be understood, just respond "I am sorry I could not understand. Could you please try asking again?"
    ---------------
    Context: ```{context}```
    '''

    user_template = '''
    Question: ```{question}```
    '''

    messages= [
    SystemMessagePromptTemplate.from_template(system_template),
    HumanMessagePromptTemplate.from_template(user_template)
    ]

    qa_prompt = ChatPromptTemplate.from_messages(messages)
    print(f"QA Prompt: {qa_prompt}")

    chain = ConversationalRetrievalChain.from_llm(llm=model, chain_type='stuff',
                                                  retriever=vector_store.as_retriever(search_kwargs={"k": 2}),
                                                 memory=memory,
                                                 combine_docs_chain_kwargs={'prompt': qa_prompt },)

    # prompt = PromptTemplate(template = prompt_template, input_variables = ["chat_history", "question"])
    qa_chain = chain.run(chat_history=chat_history, question=user_question)
    
    return qa_chain

def process_user_question(user_question,conversation_history):
    vector_store = Pinecone.from_existing_index(index_name, init_embeddings())
    # Flatten the conversation history into a single string
    context = " ".join(conversation_history)
    print(f"Context for this query: {user_question} is {context}")
    response = get_conversational_chain(vector_store,user_question,context)

    #print(response)
    st.write("Reply: ", response)
    return response

def openai_text_to_speech(text, filename="speech.mp3"):
    try:
        client = OpenAI()
        response = client.audio.speech.create(
            model='tts-1',
            input=text,
            voice='nova',
            response_format="mp3"
        )
        # If the API returns a direct binary content (audio file)
        if response.content:
            with open(filename, "wb") as out:
                out.write(response.content)
                print(f"Audio content written to file {filename}")
        # If the API returns a JSON with a URL to the audio file (not directly supported as per my last update, hypothetical usage)
        elif 'data' in response and 'url' in response['data']:
            audio_url = response["data"]["url"]
            # Here, you'd need to download the audio file from the URL, which depends on the API's actual response structure
        
        return filename
    except Exception as e:
        print(f"An error occurred during TTS operation: {e}")
        return None

def convert_audio_to_base64(audio_file_path):
    with open(audio_file_path, "rb") as audio_file:
        return base64.b64encode(audio_file.read()).decode('utf-8')

def main():
    # Call to initialize session state variables
    initialize_session_state()

    # Initialize session state for conversation history
    if 'conversation_history' not in st.session_state:
        st.session_state['conversation_history'] = []

    st.set_page_config("Chat AI")
    st.header("Chat with PDF using OpenAI")
    for i in range(0, len(st.session_state['conversation_history']), 2):
        with st.container():
            query = st.session_state['conversation_history'][i]
            response = st.session_state['conversation_history'][i + 1] if i + 1 < len(st.session_state['conversation_history']) else ""
            combined_text = query + "\n\n" + response
            st.text_area("Conversation", combined_text, height=150, disabled=True)

    user_question = st.text_input("Ask a Question from the PDF Files")

    with st.sidebar:
        st.title("Menu:")
        # File uploader for multiple types
        uploaded_files = st.file_uploader("Upload your document files", accept_multiple_files=True, type=['pdf', 'txt', 'md', 'docx'])
        
        # Input for URL
        url = st.text_input("Or enter a URL of a webpage")

        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                all_text = ""
                # Process uploaded files
                for uploaded_file in uploaded_files:
                    metadata = uploaded_file.name
                    all_text += process_uploaded_file(uploaded_file)
                    

                # Process URL
                if url:
                    all_text += scrape_webpage(url)
                    metadata=url
                #raw_text = get_pdf_text(pdf_docs)
                text_chunks = get_text_chunks(all_text,metadata)
                get_vector_store(text_chunks)
                st.success("Done")


    if user_question:
        st.session_state['conversation_history'].append("You: " + user_question)
        response = process_user_question(user_question, st.session_state['conversation_history']) 
        st.session_state['conversation_history'].append("AI: " + response)
        
        # Convert response to speech
        #audio_file = text_to_speech(response)
        audio_file = openai_text_to_speech(response)        
        audio_file_path = os.path.join(os.getcwd(), audio_file)
        audio_base64 = convert_audio_to_base64(audio_file_path)
        
        # Embed audio in HTML with autoplay
        audio_html = f'<audio autoplay><source src="data:audio/mp3;base64,{audio_base64}" type="audio/mpeg"></audio>'
        st.markdown(audio_html, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
